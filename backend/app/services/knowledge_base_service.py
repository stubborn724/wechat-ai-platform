"""Knowledge base service: document parsing, chunking, embedding, vector search."""

import hashlib
import io
import logging
import re
import uuid
from datetime import datetime
from typing import List, Optional

from sqlalchemy import func, select, text
from sqlalchemy.orm import Session

from app.models.pg_models import KbDocument, KbDocumentChunk, KnowledgeBase
from app.services.embedding_service import embed_batch_sync, embed_text_sync

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

DEFAULT_CHUNK_SIZE = 512
DEFAULT_CHUNK_OVERLAP = 64


def chunk_text(text: str, chunk_size: int = DEFAULT_CHUNK_SIZE,
               overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks at paragraph/sentence boundaries.

    Preserves paragraph structure when possible.
    """
    if not text:
        return []

    # Normalize whitespace
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    chunks: List[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)

        # Try to break at paragraph boundary
        if end < text_len:
            # Look backwards for double newline within a window
            para_break = text.rfind('\n\n', start, end)
            if para_break > start + chunk_size // 2:
                end = para_break
            else:
                # Look backwards for single newline
                line_break = text.rfind('\n', start, end)
                if line_break > start + chunk_size // 2:
                    end = line_break
                else:
                    # Look backwards for sentence end
                    sent_break = max(text.rfind('. ', start, end),
                                     text.rfind('。', start, end),
                                     text.rfind('!', start, end),
                                     text.rfind('! ', start, end))
                    if sent_break > start + chunk_size // 3:
                        end = sent_break + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Advance with overlap
        next_start = end - overlap
        if next_start <= start:
            next_start = end
        start = next_start

        # Safety: ensure progress
        if start >= text_len:
            break

    return chunks


# ---------------------------------------------------------------------------
# Document parsing
# ---------------------------------------------------------------------------


def parse_document(file_bytes: bytes, filename: str) -> Optional[str]:
    """Parse an uploaded document and return its plain text content.

    Supports PDF, DOCX, MD, and TXT files.
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    try:
        if ext == 'pdf':
            return _parse_pdf(file_bytes)
        elif ext == 'docx':
            return _parse_docx(file_bytes)
        elif ext in ('md', 'markdown'):
            return file_bytes.decode('utf-8', errors='replace')
        elif ext in ('txt', 'text'):
            return file_bytes.decode('utf-8', errors='replace')
        else:
            # Try as plain text
            return file_bytes.decode('utf-8', errors='replace')
    except Exception as exc:
        logger.error("Failed to parse %s: %s", filename, exc)
        return None


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfminer."""
    from pdfminer.high_level import extract_text
    from pdfminer.pdfparser import PDFSyntaxError

    try:
        text = extract_text(io.BytesIO(file_bytes))
        return text.strip()
    except PDFSyntaxError as exc:
        logger.warning("PDF parse error, falling back to plain text: %s", exc)
        return file_bytes.decode('utf-8', errors='replace').strip()


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n\n'.join(paragraphs)


# ---------------------------------------------------------------------------
# CRUD
# ---------------------------------------------------------------------------


def create_knowledge_base(db: Session, tenant_id: int, name: str,
                          kb_type: str = "article",
                          description: Optional[str] = None) -> KnowledgeBase:
    """Create a new knowledge base."""
    slug = _slugify(name)
    # Ensure unique slug within tenant
    base_slug = slug
    counter = 1
    while db.query(KnowledgeBase).filter(
            KnowledgeBase.slug == slug).first():
        slug = f"{base_slug}-{counter}"
        counter += 1

    kb = KnowledgeBase(
        tenant_id=tenant_id,
        name=name,
        slug=slug,
        kb_type=kb_type,
        description=description or "",
    )
    db.add(kb)
    db.commit()
    db.refresh(kb)
    logger.info("Created knowledge base: %s (slug=%s, tenant=%d)", name, slug, tenant_id)
    return kb


def list_knowledge_bases(db: Session, tenant_id: int) -> List[KnowledgeBase]:
    """List all knowledge bases for a tenant."""
    return (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.tenant_id == tenant_id,
                KnowledgeBase.is_active == 1)
        .order_by(KnowledgeBase.id.desc())
        .all()
    )


def get_knowledge_base(db: Session, kb_id: int, tenant_id: Optional[int] = None) -> Optional[KnowledgeBase]:
    """Get a single knowledge base by id."""
    query = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id)
    if tenant_id is not None:
        query = query.filter(KnowledgeBase.tenant_id == tenant_id)
    return query.first()


def delete_knowledge_base(db: Session, kb_id: int, tenant_id: Optional[int] = None) -> bool:
    """Soft-delete a knowledge base (set is_active=0)."""
    query = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id)
    if tenant_id is not None:
        query = query.filter(KnowledgeBase.tenant_id == tenant_id)
    kb = query.first()
    if not kb:
        return False
    kb.is_active = 0
    # Also soft-delete documents
    db.query(KbDocument).filter(
        KbDocument.knowledge_base_id == kb_id
    ).update({"status": "deleted"}, synchronize_session=False)
    db.commit()
    logger.info("Deleted knowledge base: id=%d", kb_id)
    return True


# ---------------------------------------------------------------------------
# Document management
# ---------------------------------------------------------------------------


def process_document(db: Session, kb_id: int, tenant_id: int,
                     file_bytes: bytes, filename: str) -> KbDocument:
    """Upload, parse, chunk, embed, and store a document.

    Returns the KbDocument row (status=ready on success, status=error on failure).
    """
    doc = KbDocument(
        knowledge_base_id=kb_id,
        filename=filename,
        file_type=filename.lower().rsplit('.', 1)[-1] if '.' in filename else "unknown",
        file_size=len(file_bytes),
        status="processing",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    try:
        # 1. Parse
        raw_text = parse_document(file_bytes, filename)
        if not raw_text or not raw_text.strip():
            doc.status = "error"
            doc.error_message = "No extractable text found"
            db.commit()
            return doc

        # 2. Dedup check via content hash
        content_hash = hashlib.sha256(raw_text.encode('utf-8')).hexdigest()
        doc.content_hash = content_hash

        existing = db.query(KbDocument).filter(
            KbDocument.knowledge_base_id == kb_id,
            KbDocument.content_hash == content_hash,
            KbDocument.status == "ready",
        ).first()
        if existing:
            logger.info("Duplicate document detected: %s (existing doc_id=%d)",
                        filename, existing.id)
            doc.status = "duplicate"
            doc.error_message = f"Duplicate of document id={existing.id}"
            db.commit()
            return doc

        # 3. Chunk
        chunks = chunk_text(raw_text)
        if not chunks:
            doc.status = "error"
            doc.error_message = "Text is empty after parsing"
            db.commit()
            return doc

        # 4. Embed batches (32 at a time to avoid token limits)
        all_vectors: List[List[float]] = []
        batch_size = 32
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            vectors = embed_batch_sync(batch)
            all_vectors.extend(vectors)
            logger.debug("Embedded batch %d/%d (%d chunks)",
                         i // batch_size + 1, (len(chunks) + batch_size - 1) // batch_size,
                         len(batch))

        # 5. Bulk insert chunks
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
        kb_type = kb.kb_type if kb else "article"
        chunk_rows = []
        for idx, (chunk_text_val, vector) in enumerate(zip(chunks, all_vectors)):
            chunk_rows.append({
                "tenant_id": tenant_id,
                "knowledge_base_id": kb_id,
                "document_id": doc.id,
                "chunk_index": idx,
                "content": chunk_text_val,
                "kb_type": kb_type,
                "embedding": vector,
            })

        # Bulk insert using raw SQL for performance with pgvector
        for row in chunk_rows:
            chunk = KbDocumentChunk(
                tenant_id=row["tenant_id"],
                knowledge_base_id=row["knowledge_base_id"],
                document_id=row["document_id"],
                chunk_index=row["chunk_index"],
                content=row["content"],
                kb_type=row["kb_type"],
                embedding=row["embedding"],
            )
            db.add(chunk)
        db.flush()

        # 6. Update document status
        doc.status = "ready"
        doc.chunk_count = len(chunks)
        db.commit()
        db.refresh(doc)
        logger.info("Document processed: %s (%d chunks, %d KB)",
                    filename, len(chunks), len(file_bytes) // 1024)
        return doc

    except Exception as exc:
        logger.error("Document processing failed: %s", exc)
        import traceback
        traceback.print_exc()
        doc.status = "error"
        doc.error_message = str(exc)[:500]
        db.commit()
        return doc


def list_documents(db: Session, kb_id: int,
                    tenant_id: Optional[int] = None) -> List[KbDocument]:
    """列出知识库文档，并通过知识库本身完成租户边界校验。

    ``KbDocument`` 不冗余保存 ``tenant_id``，必须先校验 ``KnowledgeBase``，
    不能访问不存在的文档租户字段。
    """
    if tenant_id is not None and not get_knowledge_base(db, kb_id, tenant_id):
        return []
    query = db.query(KbDocument).filter(
        KbDocument.knowledge_base_id == kb_id,
        KbDocument.status != "deleted",
    )
    return query.order_by(KbDocument.id.desc()).all()


def get_document(db: Session, doc_id: int,
                 tenant_id: Optional[int] = None) -> Optional[KbDocument]:
    """获取单份文档，并通过所属知识库而非文档字段校验租户。"""
    query = db.query(KbDocument).filter(KbDocument.id == doc_id)
    document = query.first()
    if not document or tenant_id is None:
        return document
    return document if get_knowledge_base(db, document.knowledge_base_id, tenant_id) else None


def list_document_chunks(db: Session, document_id: int) -> List[dict]:
    """返回文档的已入库切片，供前端只读预览按原文顺序展示。"""
    chunks = (
        db.query(KbDocumentChunk)
        .filter(KbDocumentChunk.document_id == document_id)
        .order_by(KbDocumentChunk.chunk_index.asc())
        .all()
    )
    return [
        {"chunk_index": chunk.chunk_index, "content": chunk.content}
        for chunk in chunks
    ]


def delete_document(db: Session, doc_id: int,
                    tenant_id: Optional[int] = None) -> bool:
    """删除文档及切片，并通过所属知识库校验租户。"""
    query = db.query(KbDocument).filter(KbDocument.id == doc_id)
    doc = query.first()
    if not doc:
        return False
    if tenant_id is not None and not get_knowledge_base(db, doc.knowledge_base_id, tenant_id):
        return False
    # Delete chunks first
    db.query(KbDocumentChunk).filter(
        KbDocumentChunk.document_id == doc_id
    ).delete()
    db.delete(doc)
    db.commit()
    return True


# ---------------------------------------------------------------------------
# Vector search
# ---------------------------------------------------------------------------


def search_knowledge_base(db: Session, kb_id: int, query: str,
                           top_k: int = 5,
                           tenant_id: Optional[int] = None) -> List[dict]:
    """Vector similarity search within a knowledge base.

    Embeds the query text, then performs cosine similarity search via pgvector
    ``<=>`` operator on KbDocumentChunk.
    """
    if not query or not query.strip():
        return []

    # 1. Embed the query
    query_vector = embed_text_sync(query)

    # 2. Search via pgvector cosine distance (<1=> operator)
    sql = text("""
        SELECT
            id, content, chunk_index,
            kb_type,
            (embedding <=> :query_vec::vector) AS distance
        FROM kb_document_chunks
        WHERE knowledge_base_id = :kb_id
          AND embedding IS NOT NULL
        ORDER BY embedding <=> :query_vec::vector
        LIMIT :top_k
    """)

    rows = db.execute(sql, {
        "query_vec": str(query_vector),
        "kb_id": kb_id,
        "top_k": top_k,
    }).fetchall()

    results = []
    for row in rows:
        results.append({
            "id": row[0],
            "content": row[1],
            "chunk_index": row[2],
            "kb_type": row[3],
            "score": float(1 - row[4]),  # cosine distance -> similarity score
        })

    return results


def search_all_knowledge_bases(db: Session, tenant_id: int, query: str,
                                top_k: int = 5) -> List[dict]:
    """Search across all active knowledge bases for a tenant."""
    if not query or not query.strip():
        return []

    query_vector = embed_text_sync(query)

    sql = text("""
        SELECT
            c.id, c.content, c.chunk_index, c.kb_type,
            c.knowledge_base_id,
            (c.embedding <=> :query_vec::vector) AS distance
        FROM kb_document_chunks c
        JOIN knowledge_bases kb ON kb.id = c.knowledge_base_id
        WHERE kb.tenant_id = :tenant_id
          AND kb.is_active = 1
          AND c.embedding IS NOT NULL
        ORDER BY c.embedding <=> :query_vec::vector
        LIMIT :top_k
    """)

    rows = db.execute(sql, {
        "query_vec": str(query_vector),
        "tenant_id": tenant_id,
        "top_k": top_k,
    }).fetchall()

    results = []
    for row in rows:
        results.append({
            "id": row[0],
            "content": row[1],
            "chunk_index": row[2],
            "kb_type": row[3],
            "knowledge_base_id": row[4],
            "score": float(1 - row[5]),
        })

    return results


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _slugify(name: str) -> str:
    """Convert a name to a URL-friendly slug."""
    slug = name.lower()
    slug = re.sub(r'[^\w\s-]', '', slug)
    slug = re.sub(r'[\s_]+', '-', slug)
    slug = re.sub(r'-+', '-', slug)
    return slug.strip('-')[:128] or "kb"
